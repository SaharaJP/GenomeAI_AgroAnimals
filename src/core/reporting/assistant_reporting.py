from __future__ import annotations

import json
import os
import re
from pathlib import Path
from typing import Any, Dict, List, Optional, Tuple


def _sanitize_llm_numbers(text: str, fact_pack_str: str) -> str:
    """Ensure generated text doesn't introduce numbers not present in fact_pack."""

    def repl(m: re.Match) -> str:
        tok = m.group(0)
        if tok in fact_pack_str:
            return tok
        return "n/a"

    return re.sub(r"\b\d+(?:\.\d+)?\b", repl, text)


JsonDict = Dict[str, Any]
NarrativeDict = Dict[str, str]


def generate_assistant_report_text_fallback(fact_pack: JsonDict) -> NarrativeDict:
    """Template narrative (always available)."""
    v = fact_pack["versions"]
    qc_status = fact_pack.get("qc", {}).get("qc_status", "NA")
    ml_metrics = fact_pack.get("ml", {}).get("metrics", {}) or {}
    mae = ml_metrics.get("mae", "NA")
    rmse = ml_metrics.get("rmse", "NA")
    rc = fact_pack.get("scoring", {}).get("row_counts", {}) or {}

    exec_lines = [
        f"Связка версий: data_version={v.get('data_version')}, qc_run={v.get('qc_run')}, model_version={v.get('model_version')}, scoring_run={v.get('scoring_run')}",
        f"QC статус: {qc_status}",
        f"ML метрики (holdout без утечек): MAE={mae}, RMSE={rmse}",
        f"Скоринг: животных в ранжировании={rc.get('n_animals_ranked', 'NA')}, PRIORITY={rc.get('n_priority','NA')}, OBSERVE={rc.get('n_observe','NA')}, CULL={rc.get('n_cull_candidates','NA')}",
    ]

    prod_explain = fact_pack.get("productivity_explainability", {}) or {}
    if prod_explain.get("available"):
        explain = prod_explain.get("explainability") or {}
        if explain.get("available"):
            tf_counts = explain.get("top_feature_counts") or {}
            if tf_counts:
                exec_lines.append("Explainability продуктивности: частые top-факторы = " + json.dumps(tf_counts, ensure_ascii=False))
            for item in (explain.get("counterfactuals_preview") or [])[:2]:
                exec_lines.append(f"Контрфакт продуктивности animal_id={item.get('animal_id','NA')}: {item.get('counterfactuals_text','NA')}")

    rec_lines = [
        "Рекомендации сформированы из результатов скоринга и правил decision-support.",
        "Списки PRIORITY/OBSERVE/CULL представлены в приложениях.",
        "При низком доверии (малые группы/пропуски/выбросы) решения требуют ручной проверки.",
    ]

    prod_explain = fact_pack.get("productivity_explainability", {}) or {}
    mast = fact_pack.get("mastitis_risk", {}) or {}
    if prod_explain.get("available"):
        rows = prod_explain.get("animal_explainability") or []
        if rows:
            rec_lines.append("\nExplainability продуктивности — почему животные попали в ранжирование:")
            for r in rows[:10]:
                line = f"- animal_id={r.get('animal_id','NA')} pred={r.get('prediction','NA')} confidence={r.get('confidence','NA')} action={r.get('action','NA')}"
                tf = str(r.get('explain_top_factors_text') or '').strip()
                cf = str(r.get('explain_counterfactuals_text') or '').strip()
                if tf and tf != 'insufficient_explainability_data':
                    line += f" | why={tf}"
                if cf and cf != 'no_simple_counterfactual':
                    line += f" | counterfactual={cf}"
                rec_lines.append(line)

    if mast.get("available"):
        top = mast.get("top_risk") or []
        exec_lines.append(f"Риск мастита: asof_date={mast.get('asof_date')}, horizon_days={mast.get('horizon_days')}, threshold={mast.get('risk_threshold')}")
        if top:
            top_ids = [str(x.get("animal_id") or x.get("entity_id") or x.get("row_id") or "") for x in top[:5]]
            top_ids = [x for x in top_ids if x]
            if top_ids:
                exec_lines.append("Топ-5 по риску мастита: " + ", ".join(top_ids))
        explain = mast.get("explainability") or {}
        if explain.get("available"):
            tf_counts = explain.get("top_feature_counts") or {}
            if tf_counts:
                exec_lines.append("Explainability: частые top-факторы по high-risk = " + json.dumps(tf_counts, ensure_ascii=False))
            preview = (explain.get("counterfactuals_preview") or [])[:2]
            for item in preview:
                exec_lines.append(f"Контрфакт animal_id={item.get('animal_id','NA')}: {item.get('counterfactuals_text','NA')}")

    pb_block = fact_pack.get("playbooks", {}) or {}
    pbs = pb_block.get("recommended") or []
    if pbs:
        rec_lines.append("\nРекомендуемый план действий (playbook):")
        for pb in pbs[:3]:
            name = str(pb.get("name") or "NA")
            kind = str(pb.get("target_kind") or "NA")
            tp = str(pb.get("target_type") or "NA")
            vid = str(pb.get("version_id") or "NA")
            src = str(pb.get("source") or "NA")
            fid = str(pb.get("farm_id") or "")
            rec_lines.append(f"- {name} | target={kind}:{tp} | farm_id={fid or 'GLOBAL'} | version_id={vid} | source={src}")
            steps = list(pb.get("steps") or [])
            for i, st in enumerate(steps[:8], start=1):
                title = str(st.get("title") or st.get("key") or "step")
                details = str(st.get("details") or "").strip()
                if details:
                    rec_lines.append(f"  {i}. {title}: {details}")
                else:
                    rec_lines.append(f"  {i}. {title}")

    if mast.get("available"):
        top = mast.get("top_risk") or []
        if top:
            rec_lines.append("\nРиск мастита — рекомендованные действия:")
            for r in top[:10]:
                risk_value = r.get('risk', r.get('risk_score', r.get('risk_proba', 'NA')))
                line = f"- animal_id={r.get('animal_id','NA')} risk={risk_value} confidence={r.get('confidence','NA')} action={r.get('recommended_action','NA')}"
                tf = str(r.get('explain_top_factors_text') or '').strip()
                cf = str(r.get('explain_counterfactuals_text') or '').strip()
                if tf and tf != 'insufficient_explainability_data':
                    line += f" | why={tf}"
                if cf and cf != 'no_simple_counterfactual':
                    line += f" | counterfactual={cf}"
                rec_lines.append(line)

    limitations = fact_pack.get("ml", {}).get("limitations", {}) or {}
    lim_lines = [
        "Ограничения:",
        f"- age_at_calving доступен: {bool(limitations.get('age_at_calving_available'))}",
    ]
    if limitations.get("age_at_calving_reason"):
        lim_lines.append(f"- age_at_calving примечание: {limitations['age_at_calving_reason']}")
    lim_lines.append("- Текст отчёта построен строго из fact pack; отсутствующие значения помечены как NA.")

    return {
        "executive_summary": "\n".join(exec_lines),
        "recommendations": "\n".join(rec_lines),
        "limitations": "\n".join(lim_lines),
    }


def generate_assistant_report_text_llm(
    fact_pack: JsonDict,
    *,
    model: Optional[str] = None,
    temperature: float = 0.2,
) -> Tuple[NarrativeDict, bool, Optional[str]]:
    """Optional LLM narrative generation.

    If LLM is unavailable or errors, caller should fall back to template.
    """
    try:
        from openai import OpenAI  # type: ignore
    except Exception:
        return {}, False, "openai_python_package_not_installed"

    api_key = os.getenv("OPENAI_API_KEY")
    if not api_key:
        return {}, False, "OPENAI_API_KEY_not_set"

    client = OpenAI(api_key=api_key)
    fact_pack_str = json.dumps(fact_pack, ensure_ascii=False, indent=2)
    chosen_model = model or os.getenv("GENOMEAI_OPENAI_MODEL", "gpt-4o-mini")

    system = (
        "Вы — ассистент по аналитическому отчёту в животноводстве. "
        "Вы ДОЛЖНЫ использовать ТОЛЬКО факты из предоставленного JSON fact_pack. "
        "Запрещено придумывать числа/проценты/суммы/датировки. "
        "Если факта нет — пишите 'NA'. "
        "Выведите 3 секции: Executive summary, Recommendations, Limitations. "
        "Если в fact_pack есть playbooks.recommended — включите их как 'рекомендуемый план действий' в секции Recommendations (кратко, по шагам). "
        "Не добавляйте никаких таблиц — только текст."
    )
    user = f"fact_pack_json:\n{fact_pack_str}"

    try:
        resp = client.chat.completions.create(
            model=chosen_model,
            messages=[
                {"role": "system", "content": system},
                {"role": "user", "content": user},
            ],
            temperature=temperature,
        )
        text = resp.choices[0].message.content or ""
    except Exception as e:
        return {}, False, f"llm_error:{type(e).__name__}"

    text = _sanitize_llm_numbers(text, fact_pack_str)

    sections: NarrativeDict = {"executive_summary": "", "recommendations": "", "limitations": ""}
    cur: Optional[str] = None
    for line in text.splitlines():
        l = line.strip()
        low = l.lower()
        if "executive" in low and "summary" in low:
            cur = "executive_summary"
            continue
        if "recommend" in low:
            cur = "recommendations"
            continue
        if "limit" in low:
            cur = "limitations"
            continue
        if cur:
            sections[cur] += (line + "\n")

    if not any(sections.values()):
        sections["executive_summary"] = text
    sections = {k: v.strip() for k, v in sections.items()}
    return sections, True, None


def render_assistant_report_docx(
    *,
    fact_pack: JsonDict,
    narrative: NarrativeDict,
    out_path: Path,
    report_version: str,
    llm_used: bool,
) -> None:
    from docx import Document  # type: ignore

    v = fact_pack["versions"]
    qc_status = fact_pack.get("qc", {}).get("qc_status", "NA")
    ml_metrics = fact_pack.get("ml", {}).get("metrics", {}) or {}
    scoring_counts = fact_pack.get("scoring", {}).get("row_counts", {}) or {}

    doc = Document()
    doc.add_heading(f"GenomeAI AgroAnimals — Отчёт {report_version}", level=0)
    doc.add_paragraph(f"Сформирован (UTC): {fact_pack.get('created_at_utc')}")
    doc.add_paragraph(f"Режим: {'LLM' if llm_used else 'Fallback'} (источник правды: fact_pack.json)")

    doc.add_heading("Связка версий", level=1)
    t = doc.add_table(rows=1, cols=2)
    hdr = t.rows[0].cells
    hdr[0].text = "Ключ"
    hdr[1].text = "Значение"
    for k in ["data_version", "qc_run", "model_version", "scoring_run"]:
        row = t.add_row().cells
        row[0].text = k
        row[1].text = str(v.get(k, "NA"))

    doc.add_heading("Executive summary", level=1)
    doc.add_paragraph(narrative.get("executive_summary", "NA"))

    doc.add_heading("QC итоги", level=1)
    doc.add_paragraph(f"QC статус: {qc_status}")
    qc_metrics = fact_pack.get("qc", {}).get("metrics", {}) or {}
    if qc_metrics:
        doc.add_paragraph("Ключевые метрики QC (агрегаты):")
        for k, val in list(qc_metrics.items())[:20]:
            doc.add_paragraph(f"- {k}: {val}")
    else:
        doc.add_paragraph("QC метрики: NA")

    doc.add_heading("ML итоги", level=1)
    doc.add_paragraph(f"Метрики: MAE={ml_metrics.get('mae','NA')}, RMSE={ml_metrics.get('rmse','NA')}")
    doc.add_paragraph(f"Цель: {fact_pack.get('ml',{}).get('target','NA')}")
    doc.add_paragraph("Признаки:")
    feats = fact_pack.get("ml", {}).get("features", {}) or {}
    doc.add_paragraph(f"- numeric: {', '.join(feats.get('numeric', []) or []) or 'NA'}")
    doc.add_paragraph(f"- categorical: {', '.join(feats.get('categorical', []) or []) or 'NA'}")

    doc.add_heading("Recommendations", level=1)
    doc.add_paragraph(narrative.get("recommendations", "NA"))

    prod_explain = fact_pack.get("productivity_explainability", {}) or {}
    mast = fact_pack.get("mastitis_risk", {}) or {}
    doc.add_heading("Explainability (productivity ML)", level=1)
    explain_prod = prod_explain.get("explainability") or {}
    if explain_prod.get("available"):
        tf_counts = explain_prod.get("top_feature_counts") or {}
        if tf_counts:
            doc.add_paragraph("Частые top-факторы по продуктивности: " + json.dumps(tf_counts, ensure_ascii=False))
        for item in (explain_prod.get("top_factors_preview") or [])[:5]:
            doc.add_paragraph(f"- animal_id={item.get('animal_id','NA')}: {item.get('top_factors_text','NA')}")
        for item in (explain_prod.get("counterfactuals_preview") or [])[:3]:
            doc.add_paragraph(f"- Контрфакт animal_id={item.get('animal_id','NA')}: {item.get('counterfactuals_text','NA')}")
    else:
        doc.add_paragraph("NA")

    doc.add_heading("Explainability (mastitis risk)", level=1)
    explain = mast.get("explainability") or {}
    if explain.get("available"):
        tf_counts = explain.get("top_feature_counts") or {}
        if tf_counts:
            doc.add_paragraph("Частые top-факторы у high-risk животных: " + json.dumps(tf_counts, ensure_ascii=False))
        for item in (explain.get("top_factors_preview") or [])[:5]:
            doc.add_paragraph(f"- animal_id={item.get('animal_id','NA')}: {item.get('top_factors_text','NA')}")
        for item in (explain.get("counterfactuals_preview") or [])[:3]:
            doc.add_paragraph(f"- Контрфакт animal_id={item.get('animal_id','NA')}: {item.get('counterfactuals_text','NA')}")
    else:
        doc.add_paragraph("NA")

    doc.add_heading("Сводка скоринга", level=1)
    doc.add_paragraph(
        f"Животных в ранжировании={scoring_counts.get('n_animals_ranked','NA')}; "
        f"PRIORITY={scoring_counts.get('n_priority','NA')}; OBSERVE={scoring_counts.get('n_observe','NA')}; "
        f"CULL={scoring_counts.get('n_cull_candidates','NA')}."
    )

    doc.add_heading("Limitations", level=1)
    doc.add_paragraph(narrative.get("limitations", "NA"))

    doc.add_page_break()
    doc.add_heading("Приложения", level=1)

    def add_table(title: str, rows: List[JsonDict]) -> None:
        doc.add_heading(title, level=2)
        if not rows:
            doc.add_paragraph("NA")
            return
        cols = list(rows[0].keys())
        tbl = doc.add_table(rows=1, cols=len(cols))
        for j, c in enumerate(cols):
            tbl.rows[0].cells[j].text = c
        for r in rows:
            rr = tbl.add_row().cells
            for j, c in enumerate(cols):
                rr[j].text = str(r.get(c, ""))

    add_table("TOP PRIORITY (до 20)", fact_pack.get("top_lists", {}).get("priority", []))
    add_table("TOP OBSERVE (до 20)", fact_pack.get("top_lists", {}).get("observe", []))
    add_table("TOP CULL CANDIDATES (до 20)", fact_pack.get("top_lists", {}).get("cull_candidates", []))

    doc.add_heading("Распределения", level=2)
    dist = fact_pack.get("distributions", {}) or {}
    if dist:
        for k, d in dist.items():
            doc.add_paragraph(f"{k}: {json.dumps(d, ensure_ascii=False)}")
    else:
        doc.add_paragraph("NA")

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(out_path)


def render_assistant_report_pdf(
    *,
    narrative: NarrativeDict,
    fact_pack: JsonDict,
    out_path: Path,
    report_version: str,
    llm_used: bool,
) -> bool:
    try:
        from reportlab.lib.pagesizes import A4  # type: ignore
        from reportlab.lib.units import cm  # type: ignore
        from reportlab.pdfbase import pdfmetrics  # type: ignore
        from reportlab.pdfbase.ttfonts import TTFont  # type: ignore
        from reportlab.pdfgen import canvas  # type: ignore
    except Exception:
        return False

    try:
        pdfmetrics.registerFont(TTFont("DejaVu", "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf"))
        font_name = "DejaVu"
    except Exception:
        font_name = "Helvetica"

    out_path.parent.mkdir(parents=True, exist_ok=True)
    c = canvas.Canvas(str(out_path), pagesize=A4)
    _w, h = A4
    x = 2 * cm
    y = h - 2 * cm
    c.setFont(font_name, 14)
    c.drawString(x, y, f"GenomeAI AgroAnimals — Report {report_version}")
    y -= 1 * cm
    c.setFont(font_name, 10)
    c.drawString(x, y, f"UTC: {fact_pack.get('created_at_utc')} | Mode: {'LLM' if llm_used else 'Fallback'}")
    y -= 1 * cm

    def para(title: str, text: str) -> None:
        nonlocal y
        c.setFont(font_name, 12)
        c.drawString(x, y, title)
        y -= 0.6 * cm
        c.setFont(font_name, 10)
        for line in (text or "NA").splitlines():
            if y < 2 * cm:
                c.showPage()
                y = h - 2 * cm
                c.setFont(font_name, 10)
            c.drawString(x, y, line[:120])
            y -= 0.45 * cm
        y -= 0.4 * cm

    para("Executive summary", narrative.get("executive_summary", "NA"))
    para("Recommendations", narrative.get("recommendations", "NA"))
    para("Limitations", narrative.get("limitations", "NA"))

    c.save()
    return True
